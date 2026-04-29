"""MVR 검토 엔진.

입력: .docx 파일 + 룰셋 YAML
출력: 위반 사항 리스트 (룰 ID, 위치, severity, 메시지)

지원 check_type:
  - regex_replace : 본문에서 잘못된 표현 찾기
  - regex_find    : 본문에서 패턴 찾기 (forbidden 여부에 따라)
  - table_header  : 특정 표의 헤더 검증
  - required_section : 섹션/키워드 존재 확인
  - format_check  : 폰트/행간 등 서식 검증
"""

import re
import yaml
from dataclasses import dataclass, asdict
from pathlib import Path
from docx import Document
from docx.shared import Pt


@dataclass
class Violation:
    rule_id: str
    rule_name: str
    severity: str  # required / recommended / info
    location: str  # "본문 단락 12", "표 2", "전역" 등
    message: str
    matched_text: str = ""  # 실제로 발견된 텍스트


# ============================================================
# 문서 로더
# ============================================================

def load_document(path):
    """docx 파일 읽어서 단락/표/서식 정보 추출."""
    doc = Document(path)

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            # 첫 run의 폰트 크기 가져오기
            font_size = None
            for run in para.runs:
                if run.font.size:
                    font_size = run.font.size.pt
                    break
            paragraphs.append({
                "index": i,
                "text": para.text,
                "font_size": font_size,
            })

    tables = []
    for i, table in enumerate(doc.tables, start=1):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        body_rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows[1:]
        ]
        tables.append({
            "index": i,  # 표 번호 (1부터)
            "headers": headers,
            "rows": body_rows,
        })

    full_text = "\n".join(p["text"] for p in paragraphs)

    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": full_text,
    }


# ============================================================
# 룰 적용 함수들 (check_type별)
# ============================================================

def check_regex_replace(rule, doc_data):
    """잘못된 표현 → 올바른 표현 룰. forbidden 패턴 발견 시 위반."""
    violations = []
    pattern = rule["pattern"]
    for p in doc_data["paragraphs"]:
        for m in re.finditer(pattern, p["text"]):
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location=f"본문 단락 {p['index']}",
                message=rule["message"],
                matched_text=m.group(0),
            ))
    return violations


def check_regex_find(rule, doc_data):
    """패턴 찾기 룰. forbidden=True면 발견 시 flag, False면 미발견 시 flag."""
    violations = []
    pattern = rule["pattern"]
    forbidden = rule.get("forbidden", True)

    found_locations = []
    # 본문에서 찾기
    for p in doc_data["paragraphs"]:
        for m in re.finditer(pattern, p["text"]):
            found_locations.append((f"본문 단락 {p['index']}", m.group(0)))

    # 표에서도 찾기 (헤더 + 본문)
    for t in doc_data["tables"]:
        for h in t["headers"]:
            if re.search(pattern, h):
                found_locations.append((f"표 {t['index']} 헤더", h))
        for row_i, row in enumerate(t["rows"]):
            for cell in row:
                if re.search(pattern, cell):
                    found_locations.append((f"표 {t['index']} 본문", cell))

    if forbidden:
        for loc, txt in found_locations:
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location=loc,
                message=rule["message"],
                matched_text=txt,
            ))
    else:
        if not found_locations:
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location="전역",
                message=rule["message"],
            ))

    return violations


def check_table_header(rule, doc_data):
    """특정 표의 헤더 검증."""
    violations = []
    target_idx = rule["table_index"]

    target_table = next(
        (t for t in doc_data["tables"] if t["index"] == target_idx),
        None,
    )
    if not target_table:
        violations.append(Violation(
            rule_id=rule["id"],
            rule_name=rule["name"],
            severity=rule["severity"],
            location=f"표 {target_idx}",
            message=f"표 {target_idx}을(를) 찾지 못함",
        ))
        return violations

    headers = target_table["headers"]
    headers_joined = " | ".join(headers)

    # 금지 키워드 확인
    forbidden = rule.get("forbidden_keywords", [])
    for kw in forbidden:
        if any(kw in h for h in headers):
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location=f"표 {target_idx} 헤더",
                message=rule["message"],
                matched_text=f'금지어 "{kw}" 포함: {headers_joined}',
            ))

    # 기대 헤더 확인 (있으면 모두 포함되어야 함)
    expected = rule.get("expected_headers", [])
    if expected:
        missing = [h for h in expected if not any(h in actual for actual in headers)]
        if missing:
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location=f"표 {target_idx} 헤더",
                message=rule["message"],
                matched_text=f"누락 헤더: {missing} / 현재: {headers_joined}",
            ))

    # 언어 검증 (en이면 한글 포함 시 위반)
    language = rule.get("language")
    if language == "en":
        for h in headers:
            if re.search(r"[가-힣]", h):
                violations.append(Violation(
                    rule_id=rule["id"],
                    rule_name=rule["name"],
                    severity=rule["severity"],
                    location=f"표 {target_idx} 헤더",
                    message=rule["message"],
                    matched_text=f'한글 포함: "{h}"',
                ))
                break

    return violations


def check_required_section(rule, doc_data):
    """필수 섹션/키워드 존재 확인."""
    violations = []
    keywords = rule.get("keywords", [])
    text = doc_data["full_text"]

    # 약어 정의 단락 룰 (예: CE, CPA, LMW)
    # 약어가 본문에 등장하는데 풀네임 정의가 없으면 위반
    section_name = rule.get("section_name", "")

    if "정의 단락" in section_name:
        # 첫 키워드가 약어, 둘째가 풀네임
        if len(keywords) >= 2:
            abbrev, fullname = keywords[0], keywords[1]
            uses_abbrev = abbrev in text
            has_fullname = fullname.lower() in text.lower()
            if uses_abbrev and not has_fullname:
                violations.append(Violation(
                    rule_id=rule["id"],
                    rule_name=rule["name"],
                    severity=rule["severity"],
                    location="전역",
                    message=rule["message"],
                    matched_text=f'"{abbrev}" 사용되나 "{fullname}" 정의 없음',
                ))
        else:
            # 키워드 하나뿐 (예: CPA만)
            abbrev = keywords[0]
            if abbrev in text:
                # 풀네임이 같이 나오는지 (괄호 형태로) 확인
                pattern = rf"{abbrev}\s*\([A-Za-z\s]+\)|\([A-Za-z\s]+\)\s*{abbrev}"
                if not re.search(pattern, text):
                    violations.append(Violation(
                        rule_id=rule["id"],
                        rule_name=rule["name"],
                        severity=rule["severity"],
                        location="전역",
                        message=rule["message"],
                        matched_text=f'"{abbrev}" 사용되나 풀네임 정의 없음',
                    ))
    else:
        # 일반 섹션 (Revision History, 서명란 등)
        # 키워드 중 하나라도 본문에 있으면 OK
        if not any(kw.lower() in text.lower() for kw in keywords):
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location="전역",
                message=rule["message"],
                matched_text=f"검색 키워드: {keywords}",
            ))

    return violations


def check_format(rule, doc_data):
    """폰트/행간 등 서식 검증 (간이)."""
    violations = []
    target = rule["target"]
    expected = rule["expected_value"]

    if target == "font_size":
        # 본문 단락 폰트 검증
        wrong_paragraphs = []
        for p in doc_data["paragraphs"]:
            if p["font_size"] is not None and p["font_size"] != expected:
                wrong_paragraphs.append((p["index"], p["font_size"]))

        if wrong_paragraphs:
            sample = wrong_paragraphs[0]
            violations.append(Violation(
                rule_id=rule["id"],
                rule_name=rule["name"],
                severity=rule["severity"],
                location=f"본문 ({len(wrong_paragraphs)}개 단락)",
                message=rule["message"],
                matched_text=f"발견: {sample[1]}pt (기대: {expected}pt)",
            ))
    else:
        # line_spacing, paragraph_spacing, alignment는
        # python-docx에서 정확한 검증이 까다로움 — v2로 미룸
        violations.append(Violation(
            rule_id=rule["id"],
            rule_name=rule["name"],
            severity="info",
            location="전역",
            message=f"[수동확인] {rule['message']}",
        ))

    return violations


# ============================================================
# 메인 검토 함수
# ============================================================

CHECKERS = {
    "regex_replace": check_regex_replace,
    "regex_find": check_regex_find,
    "table_header": check_table_header,
    "required_section": check_required_section,
    "format_check": check_format,
}


def review_document(docx_path, rules_path):
    """문서 검토 메인 함수."""
    doc_data = load_document(docx_path)

    with open(rules_path, encoding="utf-8") as f:
        rules = yaml.safe_load(f)["rules"]

    all_violations = []
    for rule in rules:
        if not rule.get("enabled", True):
            continue
        checker = CHECKERS.get(rule["check_type"])
        if not checker:
            continue
        try:
            violations = checker(rule, doc_data)
            all_violations.extend(violations)
        except Exception as e:
            print(f"[ERROR] {rule['id']} 검사 실패: {e}")

    return all_violations, doc_data


# ============================================================
# 실행
# ============================================================

if __name__ == "__main__":
    base = Path(__file__).parent
    docx_path = base / "sample_MVR_with_violations.docx"
    rules_path = base / "MVR_rules.yaml"

    violations, doc_data = review_document(docx_path, rules_path)

    print("=" * 60)
    print(f"검토 대상: {docx_path.name}")
    print(f"단락 수: {len(doc_data['paragraphs'])}, 표 수: {len(doc_data['tables'])}")
    print("=" * 60)

    # severity별 카운트
    from collections import Counter
    sev_count = Counter(v.severity for v in violations)
    print(f"\n위반 총 {len(violations)}건")
    print(f"  필수 (required)   : {sev_count.get('required', 0)}")
    print(f"  권고 (recommended): {sev_count.get('recommended', 0)}")
    print(f"  참고 (info)       : {sev_count.get('info', 0)}")

    print("\n" + "-" * 60)
    print("상세 (severity 순):")
    print("-" * 60)

    sev_order = {"required": 0, "recommended": 1, "info": 2}
    sorted_v = sorted(violations, key=lambda v: (sev_order[v.severity], v.rule_id))

    for v in sorted_v:
        marker = {"required": "🔴", "recommended": "🟡", "info": "🔵"}[v.severity]
        print(f"\n{marker} [{v.rule_id}] {v.rule_name}")
        print(f"   위치: {v.location}")
        print(f"   메시지: {v.message}")
        if v.matched_text:
            text = v.matched_text[:80] + ("..." if len(v.matched_text) > 80 else "")
            print(f"   발견: {text}")
