"""가짜 MVR 문서 생성기.

각 룰에 대해 일부러 위반 사례를 심어둠.
룰 검토 엔진이 위반을 정확히 잡는지 검증하기 위함.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    return table


def build_fake_mvr():
    doc = Document()

    # ===== 표지 =====
    title = doc.add_heading("Method Validation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Document No. MVR-BGM-NaCl-001 / Rev. 02")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ===== 1. Introduction =====
    doc.add_heading("1. Introduction", level=1)
    # RULE-MVR-201 위반: CE 약어를 풀어쓴 정의 단락 없이 사용
    # RULE-MVR-202 위반: CPA 정의 단락 없음
    # RULE-MVR-203 위반: LMW 정의 단락 없음
    doc.add_paragraph(
        "This report describes the validation of CE method for the analysis of "
        "BGM-NaCl. The CPA values and LMW species are reported. "
        "All tests were performed according to internal SOPs."
    )

    # RULE-MVR-001 위반: "허용 기준" (띄어쓰기) 사용 — 여러 번
    doc.add_paragraph(
        "분석 절차는 허용 기준에 따라 수행되었다. "
        "모든 결과는 정해진 허용 기준 내에 있어야 한다."
    )

    # ===== 2. Validation Items =====
    doc.add_heading("2. Validation Items", level=1)

    # 표 1 — RULE-MVR-101: 한글만 캡션, 영문 X (정상)
    add_table(
        doc,
        headers=["항목", "내용"],
        rows=[
            ["Identification", "Capillary Electrophoresis 분석"],
            ["Range", "0.5~2.0 mg/mL"],
        ],
        caption="표 1. 분석 개요",
    )
    doc.add_paragraph()

    # 표 2 — RULE-MVR-102 위반: "Test"가 헤더에 포함됨
    add_table(
        doc,
        headers=["Test", "Characteristics", "Type", "Selection"],
        rows=[
            ["Specificity", "Selective", "Quantitative", "Required"],
            ["Linearity", "Range", "Quantitative", "Required"],
        ],
        caption="표 2. Validation parameters",
    )
    doc.add_paragraph()

    # ===== 3. Specificity =====
    doc.add_heading("3. Specificity", level=1)
    doc.add_paragraph(
        "특이성 시험은 허용 기준 95% 이상으로 설정되었다."
    )

    # RULE-MVR-104 위반: 분석순서 표 헤더에 "Validation Item" 사용
    add_table(
        doc,
        headers=["Validation Item", "Acceptance Criteria", "Result"],
        rows=[
            ["Specificity", "≥ 95%", "98.7%"],
            ["Accuracy", "95~105%", "101.2%"],
        ],
        caption="표 3. 분석순서",
    )
    doc.add_paragraph()

    # 표 24 — RULE-MVR-103 위반: 한글 병기됨 (영문만 와야 함)
    # 표 24를 만들기 위해 더미 표들 추가
    for i in range(4, 24):
        doc.add_paragraph(f"(표 {i} 자리표시자)")

    add_table(
        doc,
        headers=["Parameters / 파라미터", "Acceptance Criteria / 허용기준", "Formula / 수식"],
        rows=[
            ["RSD", "≤ 2.0%", "SD/mean × 100"],
            ["Recovery", "98~102%", "(found/added) × 100"],
        ],
        caption="표 24. Calculation parameters",
    )
    doc.add_paragraph()

    # ===== 4. Equipment =====
    doc.add_heading("4. Equipment Calibration", level=1)

    # RULE-MVR-002 위반: "검교정 일자" 컬럼 존재
    # RULE-MVR-003 위반: "차기 일자" 영문 병기 없음
    add_table(
        doc,
        headers=["장비", "검교정 일자", "차기 일자"],
        rows=[
            ["HPLC-001", "2026.03.15", "2027.03.15"],
            ["UV-002", "2026.04.01", "2027.04.01"],
        ],
        caption="표 25. 장비 검교정 현황",
    )
    doc.add_paragraph()

    # ===== 5. Conclusion =====
    doc.add_heading("5. Conclusion", level=1)
    doc.add_paragraph(
        "모든 validation 항목이 허용 기준을 만족하였다."
    )

    # ===== 의도적으로 누락 =====
    # RULE-MVR-204 위반: Revision History 섹션 없음
    # RULE-MVR-205 위반: 서명란 (Prepared/Reviewed/Approved by) 없음

    # ===== 서식 위반 =====
    # RULE-MVR-301 위반: 본문 폰트 12pt (11pt가 표준)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size is None:
                run.font.size = Pt(12)  # 일부러 12pt

    return doc


if __name__ == "__main__":
    doc = build_fake_mvr()
    out_path = "/home/claude/mvr_reviewer/sample_MVR_with_violations.docx"
    doc.save(out_path)
    print(f"생성 완료: {out_path}")
    print("\n의도적으로 심은 위반:")
    print("  RULE-MVR-001 (허용 기준 띄어쓰기) — 본문 3회")
    print("  RULE-MVR-002 (검교정 일자 컬럼) — 표 25")
    print("  RULE-MVR-003 (차기 영문 병기) — 표 25")
    print("  RULE-MVR-102 (표 2 Test 포함) — 표 2")
    print("  RULE-MVR-103 (표 24 한글 병기) — 표 24")
    print("  RULE-MVR-104 (Validation Item 사용) — 표 3")
    print("  RULE-MVR-201 (CE 정의 단락 없음)")
    print("  RULE-MVR-202 (CPA 정의 단락 없음)")
    print("  RULE-MVR-203 (LMW 정의 단락 없음)")
    print("  RULE-MVR-204 (Revision History 없음)")
    print("  RULE-MVR-205 (서명란 없음)")
    print("  RULE-MVR-301 (폰트 12pt)")
